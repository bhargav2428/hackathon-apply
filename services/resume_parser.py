"""
Resume Parser Service - Extracts information from uploaded resumes
Enhanced version with better accuracy and validation
"""
import os
import re
from typing import Dict, Any, List, Optional, Tuple


def parse_resume(filepath: str) -> Dict[str, Any]:
    """
    Parse a resume file and extract relevant information
    
    Supports: PDF, DOC, DOCX
    """
    extension = filepath.rsplit('.', 1)[-1].lower()
    
    if extension == 'pdf':
        return _parse_pdf(filepath)
    elif extension in ['doc', 'docx']:
        return _parse_docx(filepath)
    else:
        return {'error': f'Unsupported file type: {extension}', 'validation': {'valid': False}}


def _parse_pdf(filepath: str) -> Dict[str, Any]:
    """Parse PDF resume with enhanced extraction"""
    try:
        import pdfplumber
        
        text = ""
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
        
        if not text.strip():
            return {'error': 'Could not extract text from PDF', 'text': '', 'validation': {'valid': False}}
        
        return _extract_info_enhanced(text)
    except ImportError:
        return {'error': 'pdfplumber not installed', 'text': '', 'validation': {'valid': False}}
    except Exception as e:
        return {'error': f'PDF parsing error: {str(e)}', 'text': '', 'validation': {'valid': False}}


def _parse_docx(filepath: str) -> Dict[str, Any]:
    """Parse DOCX resume with enhanced extraction"""
    try:
        from docx import Document
        
        doc = Document(filepath)
        
        # Get text from paragraphs
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also check tables (resumes often use tables for formatting)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        text = "\n".join(text_parts)
        
        if not text.strip():
            return {'error': 'Could not extract text from DOCX', 'text': '', 'validation': {'valid': False}}
        
        return _extract_info_enhanced(text)
    except ImportError:
        return {'error': 'python-docx not installed', 'text': '', 'validation': {'valid': False}}
    except Exception as e:
        return {'error': f'DOCX parsing error: {str(e)}', 'text': '', 'validation': {'valid': False}}


def _extract_info_enhanced(text: str) -> Dict[str, Any]:
    """Enhanced extraction with validation"""
    
    # Clean up text
    text = _clean_text(text)
    
    result = {
        'text': text,
        'raw_text': text,  # Keep original for debugging
        'name': None,
        'email': None,
        'phone': None,
        'github': None,
        'linkedin': None,
        'portfolio': None,
        'skills': [],
        'programming_languages': [],
        'frameworks': [],
        'databases': [],
        'tools': [],
        'education': [],
        'experience': [],
        'projects': [],
        'certifications': [],
        'validation': {
            'valid': False,
            'fields_found': [],
            'fields_missing': [],
            'confidence': 0,
            'warnings': []
        }
    }
    
    # Extract contact information
    result['email'] = _extract_email(text)
    result['phone'] = _extract_phone(text)
    result['github'] = _extract_github(text)
    result['linkedin'] = _extract_linkedin(text)
    result['portfolio'] = _extract_portfolio(text)
    result['name'] = _extract_name(text, result['email'])
    
    # Extract sections
    sections = _identify_sections(text)
    
    # Extract technical skills
    skills_text = sections.get('skills', '') or text
    result['programming_languages'] = _extract_programming_languages(skills_text)
    result['frameworks'] = _extract_frameworks(skills_text)
    result['databases'] = _extract_databases(skills_text)
    result['tools'] = _extract_tools(skills_text)
    result['skills'] = _extract_general_skills(skills_text)
    
    # Extract education
    result['education'] = _extract_education(sections.get('education', '') or text)
    
    # Extract experience
    result['experience'] = _extract_experience(sections.get('experience', '') or text)
    
    # Extract projects
    result['projects'] = _extract_projects(sections.get('projects', '') or text)
    
    # Validate extraction
    result['validation'] = _validate_extraction(result)
    
    # Combine all skills for easy access
    all_tech_skills = (
        result['programming_languages'] + 
        result['frameworks'] + 
        result['databases'] + 
        result['tools']
    )
    result['all_technical_skills'] = list(set(all_tech_skills))
    
    return result


def _clean_text(text: str) -> str:
    """Clean and normalize resume text"""
    # Remove multiple spaces
    text = re.sub(r' +', ' ', text)
    # Remove multiple newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Fix common OCR issues
    text = text.replace('|', 'I').replace('0', 'O') if len(text) < 100 else text
    return text.strip()


def _extract_email(text: str) -> Optional[str]:
    """Extract email with validation"""
    # More comprehensive email pattern
    patterns = [
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b',
        r'[A-Za-z0-9._%+-]+\s*@\s*[A-Za-z0-9.-]+\s*\.\s*[A-Za-z]{2,}',  # With spaces
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            email = matches[0].replace(' ', '')
            # Validate email format
            if re.match(r'^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$', email):
                return email.lower()
    return None


def _extract_phone(text: str) -> Optional[str]:
    """Extract phone number with various formats"""
    patterns = [
        r'\+?1?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}',  # US format
        r'\+?[0-9]{1,3}[-.\s]?[0-9]{2,4}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}',  # International
        r'\b[0-9]{10}\b',  # 10 digit
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text)
        if matches:
            phone = matches[0].strip()
            # Clean up phone number
            phone = re.sub(r'[^\d+]', '', phone)
            if len(phone) >= 10:
                return phone
    return None


def _extract_github(text: str) -> Optional[str]:
    """Extract GitHub URL"""
    patterns = [
        r'(?:https?://)?(?:www\.)?github\.com/([A-Za-z0-9_-]+)',
        r'github\s*[:\-]?\s*([A-Za-z0-9_-]+)',
        r'@([A-Za-z0-9_-]+)\s+(?:github|gh)',
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            username = matches[0]
            if len(username) > 2 and username.lower() not in ['com', 'www', 'http', 'https']:
                return f'https://github.com/{username}'
    return None


def _extract_linkedin(text: str) -> Optional[str]:
    """Extract LinkedIn URL"""
    patterns = [
        r'(?:https?://)?(?:www\.)?linkedin\.com/in/([A-Za-z0-9_-]+)',
        r'linkedin\s*[:\-]?\s*(?:https?://)?(?:www\.)?linkedin\.com/in/([A-Za-z0-9_-]+)',
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            return f'https://linkedin.com/in/{matches[0]}'
    return None


def _extract_portfolio(text: str) -> Optional[str]:
    """Extract portfolio/personal website URL"""
    patterns = [
        r'(?:portfolio|website|blog)\s*[:\-]?\s*(https?://[^\s]+)',
        r'(?:https?://)?(?:www\.)?([A-Za-z0-9_-]+\.(?:dev|io|me|com|tech|site)(?:/[^\s]*)?)',
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            url = matches[0]
            if not url.startswith('http'):
                url = 'https://' + url
            return url
    return None


def _extract_name(text: str, email: Optional[str] = None) -> Optional[str]:
    """Extract name using multiple strategies"""
    lines = text.split('\n')
    
    # Strategy 1: First non-empty line that looks like a name
    for line in lines[:10]:
        line = line.strip()
        if not line or len(line) > 60:
            continue
        
        # Skip lines with common resume headers
        skip_words = ['resume', 'cv', 'curriculum', 'vitae', 'contact', 'objective', 'summary', 'profile', '@', 'http', 'www']
        if any(word in line.lower() for word in skip_words):
            continue
        
        # Check if line looks like a name (2-4 capitalized words)
        words = line.split()
        if 1 <= len(words) <= 4:
            # Check if most words are capitalized and not all caps
            cap_words = sum(1 for w in words if w[0].isupper() and not w.isupper())
            if cap_words >= len(words) * 0.5:
                # Verify it's not a title/header
                if not re.match(r'^(Mr|Ms|Mrs|Dr|Prof)\b', line):
                    return line
    
    # Strategy 2: Extract from email
    if email:
        local_part = email.split('@')[0]
        # Try to parse name from email like john.doe or johndoe
        if '.' in local_part:
            parts = local_part.split('.')
            name = ' '.join(p.capitalize() for p in parts if len(p) > 1)
            return name
        elif '_' in local_part:
            parts = local_part.split('_')
            name = ' '.join(p.capitalize() for p in parts if len(p) > 1)
            return name
    
    return None


def _identify_sections(text: str) -> Dict[str, str]:
    """Identify and extract sections from resume"""
    sections = {}
    
    # Common section headers
    section_patterns = {
        'skills': r'(?:technical\s+)?skills|technologies|tech\s+stack|proficiencies',
        'experience': r'(?:work\s+)?experience|employment(?:\s+history)?|professional\s+experience',
        'education': r'education|academic(?:\s+background)?|qualifications',
        'projects': r'projects|portfolio|work\s+samples',
        'certifications': r'certifications?|certificates?|licenses?',
        'summary': r'(?:professional\s+)?summary|objective|profile|about(?:\s+me)?',
    }
    
    lines = text.split('\n')
    current_section = None
    section_content = []
    
    for i, line in enumerate(lines):
        line_clean = line.strip().lower()
        
        # Check if this line is a section header
        is_header = False
        for section_name, pattern in section_patterns.items():
            if re.match(f'^{pattern}\\s*:?\\s*$', line_clean) or \
               (len(line_clean) < 30 and re.search(pattern, line_clean)):
                # Save previous section
                if current_section and section_content:
                    sections[current_section] = '\n'.join(section_content)
                
                current_section = section_name
                section_content = []
                is_header = True
                break
        
        if not is_header and current_section:
            section_content.append(line)
    
    # Save last section
    if current_section and section_content:
        sections[current_section] = '\n'.join(section_content)
    
    return sections


def _extract_programming_languages(text: str) -> List[str]:
    """Extract programming languages with confidence"""
    languages = {
        'Python': [r'\bPython\b', r'\bpy\b'],
        'JavaScript': [r'\bJavaScript\b', r'\bJS\b'],
        'TypeScript': [r'\bTypeScript\b', r'\bTS\b'],
        'Java': [r'\bJava\b(?!\s*Script)'],
        'C++': [r'\bC\+\+\b', r'\bcpp\b'],
        'C#': [r'\bC#\b', r'\bC-Sharp\b', r'\bCSharp\b'],
        'C': [r'\bC\b(?!\+|\#)'],
        'Go': [r'\bGo\b', r'\bGolang\b'],
        'Rust': [r'\bRust\b'],
        'Ruby': [r'\bRuby\b'],
        'PHP': [r'\bPHP\b'],
        'Swift': [r'\bSwift\b'],
        'Kotlin': [r'\bKotlin\b'],
        'Scala': [r'\bScala\b'],
        'R': [r'\bR\b(?!\s+[Dd]evelop)'],
        'MATLAB': [r'\bMATLAB\b'],
        'Perl': [r'\bPerl\b'],
        'Dart': [r'\bDart\b'],
        'Lua': [r'\bLua\b'],
        'Julia': [r'\bJulia\b'],
        'Haskell': [r'\bHaskell\b'],
        'SQL': [r'\bSQL\b'],
        'Bash': [r'\bBash\b', r'\bShell\b'],
        'PowerShell': [r'\bPowerShell\b'],
        'Assembly': [r'\bAssembly\b', r'\bASM\b'],
        'Solidity': [r'\bSolidity\b'],
        'VHDL': [r'\bVHDL\b'],
        'Verilog': [r'\bVerilog\b'],
    }
    
    found = []
    for lang, patterns in languages.items():
        for pattern in patterns:
            if re.search(pattern, text, re.IGNORECASE):
                found.append(lang)
                break
    
    return sorted(list(set(found)))


def _extract_frameworks(text: str) -> List[str]:
    """Extract frameworks and libraries"""
    frameworks = [
        # Frontend
        'React', 'Angular', 'Vue.js', 'Vue', 'Next.js', 'Nuxt.js', 'Svelte', 'SvelteKit',
        'jQuery', 'Bootstrap', 'Tailwind CSS', 'Tailwind', 'Material UI', 'MUI',
        'Chakra UI', 'Ant Design', 'Styled Components', 'Emotion',
        # Backend
        'Django', 'Flask', 'FastAPI', 'Express.js', 'Express', 'NestJS', 'Koa',
        'Spring', 'Spring Boot', 'ASP.NET', '.NET Core', 'Ruby on Rails', 'Rails',
        'Laravel', 'Symfony', 'CodeIgniter', 'Gin', 'Echo', 'Fiber',
        # Mobile
        'React Native', 'Flutter', 'SwiftUI', 'Jetpack Compose', 'Xamarin',
        'Ionic', 'Cordova', 'PhoneGap', 'Expo',
        # ML/AI
        'TensorFlow', 'PyTorch', 'Keras', 'scikit-learn', 'sklearn',
        'Pandas', 'NumPy', 'OpenCV', 'Hugging Face', 'LangChain',
        'Matplotlib', 'Seaborn', 'Plotly', 'NLTK', 'spaCy',
        # Desktop
        'Electron', 'Tauri', 'Qt', 'GTK', 'wxWidgets',
        # Testing
        'Jest', 'Mocha', 'Chai', 'Cypress', 'Selenium', 'Playwright',
        'pytest', 'unittest', 'JUnit', 'TestNG',
        # Other
        'GraphQL', 'Apollo', 'Socket.io', 'RabbitMQ', 'Celery',
        'Node.js', 'Deno', 'Bun',
    ]
    
    found = []
    for framework in frameworks:
        # Escape special characters for regex
        pattern = r'\b' + re.escape(framework).replace(r'\.', r'\.?') + r'\b'
        if re.search(pattern, text, re.IGNORECASE):
            found.append(framework)
    
    return sorted(list(set(found)))


def _extract_databases(text: str) -> List[str]:
    """Extract database technologies"""
    databases = [
        'PostgreSQL', 'Postgres', 'MySQL', 'MariaDB', 'SQLite', 'SQL Server', 'MSSQL',
        'Oracle', 'Oracle DB', 'MongoDB', 'Mongoose', 'Redis', 'Memcached',
        'Elasticsearch', 'Cassandra', 'DynamoDB', 'Firebase', 'Firestore',
        'Neo4j', 'GraphDB', 'CouchDB', 'RethinkDB', 'InfluxDB', 'TimescaleDB',
        'Supabase', 'PlanetScale', 'Prisma', 'TypeORM', 'Sequelize', 'SQLAlchemy',
    ]
    
    found = []
    for db in databases:
        pattern = r'\b' + re.escape(db) + r'\b'
        if re.search(pattern, text, re.IGNORECASE):
            found.append(db)
    
    return sorted(list(set(found)))


def _extract_tools(text: str) -> List[str]:
    """Extract development tools and platforms"""
    tools = [
        # Version Control
        'Git', 'GitHub', 'GitLab', 'Bitbucket', 'SVN',
        # Cloud
        'AWS', 'Amazon Web Services', 'Azure', 'Google Cloud', 'GCP', 'Heroku',
        'Vercel', 'Netlify', 'DigitalOcean', 'Cloudflare',
        # DevOps
        'Docker', 'Kubernetes', 'K8s', 'Jenkins', 'Travis CI', 'CircleCI',
        'GitHub Actions', 'GitLab CI', 'Terraform', 'Ansible', 'Puppet',
        'Nginx', 'Apache', 'Linux', 'Ubuntu', 'CentOS',
        # Monitoring
        'Prometheus', 'Grafana', 'DataDog', 'New Relic', 'Sentry',
        # Design
        'Figma', 'Adobe XD', 'Sketch', 'InVision', 'Photoshop', 'Illustrator',
        # Other
        'Jira', 'Confluence', 'Slack', 'Notion', 'Trello', 'Asana',
        'Postman', 'Insomnia', 'VS Code', 'IntelliJ', 'PyCharm', 'WebStorm',
        'Webpack', 'Vite', 'Babel', 'ESLint', 'Prettier',
        'npm', 'yarn', 'pnpm', 'pip', 'poetry', 'conda',
    ]
    
    found = []
    for tool in tools:
        pattern = r'\b' + re.escape(tool) + r'\b'
        if re.search(pattern, text, re.IGNORECASE):
            found.append(tool)
    
    return sorted(list(set(found)))


def _extract_general_skills(text: str) -> List[str]:
    """Extract general/soft skills"""
    skills = [
        # Technical domains
        'Machine Learning', 'Deep Learning', 'Natural Language Processing', 'NLP',
        'Computer Vision', 'Data Science', 'Data Analysis', 'Data Engineering',
        'Big Data', 'ETL', 'Data Visualization', 'Business Intelligence',
        'Cloud Computing', 'Cloud Architecture', 'Microservices', 'Serverless',
        'DevOps', 'MLOps', 'DataOps', 'SRE', 'CI/CD',
        'REST API', 'API Development', 'gRPC', 'WebSocket', 'WebRTC',
        'Blockchain', 'Web3', 'Smart Contracts', 'DeFi', 'NFT',
        'AR/VR', 'Augmented Reality', 'Virtual Reality', 'XR',
        'IoT', 'Internet of Things', 'Embedded Systems',
        'Mobile Development', 'Web Development', 'Full Stack',
        'Frontend Development', 'Backend Development',
        'System Design', 'Software Architecture', 'Design Patterns',
        'Agile', 'Scrum', 'Kanban', 'SDLC',
        'UI/UX', 'User Interface', 'User Experience', 'Responsive Design',
        'Security', 'Cybersecurity', 'Penetration Testing', 'OWASP',
        'Performance Optimization', 'Scalability', 'Load Balancing',
        # Soft skills
        'Team Leadership', 'Project Management', 'Problem Solving',
        'Communication', 'Collaboration', 'Mentoring',
        'Technical Writing', 'Documentation', 'Code Review',
    ]
    
    found = []
    for skill in skills:
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text, re.IGNORECASE):
            found.append(skill)
    
    return sorted(list(set(found)))


def _extract_education(text: str) -> List[Dict[str, str]]:
    """Extract education information"""
    education = []
    
    # Common degree patterns
    degree_patterns = [
        r"(Bachelor'?s?|B\.?S\.?|B\.?A\.?|B\.?Tech\.?|B\.?E\.?)\s+(?:of\s+|in\s+)?([A-Za-z\s]+)",
        r"(Master'?s?|M\.?S\.?|M\.?A\.?|M\.?Tech\.?|M\.?E\.?|MBA)\s+(?:of\s+|in\s+)?([A-Za-z\s]+)",
        r"(Ph\.?D\.?|Doctorate)\s+(?:of\s+|in\s+)?([A-Za-z\s]+)",
        r"(Associate'?s?)\s+(?:of\s+|in\s+)?([A-Za-z\s]+)",
    ]
    
    # Year patterns
    year_pattern = r'(19|20)\d{2}'
    
    lines = text.split('\n')
    for line in lines:
        for pattern in degree_patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                degree_type = match.group(1)
                field = match.group(2).strip()[:50]  # Limit field length
                
                # Try to find year
                years = re.findall(year_pattern, line)
                year = years[-1] if years else None
                
                education.append({
                    'degree': degree_type,
                    'field': field,
                    'year': year,
                    'raw': line.strip()[:100]
                })
                break
    
    return education


def _extract_experience(text: str) -> List[Dict[str, str]]:
    """Extract work experience"""
    experience = []
    
    # Common job title patterns
    title_patterns = [
        r'(Senior|Junior|Lead|Principal|Staff)?\s*(Software|Backend|Frontend|Full[\s-]?Stack|DevOps|Data|ML|AI|Cloud|Mobile|Web)?\s*(Engineer|Developer|Architect|Scientist|Analyst|Designer|Manager|Lead)',
        r'(Intern|Internship)\s*[-–]?\s*([A-Za-z\s]+)',
    ]
    
    lines = text.split('\n')
    current_entry = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check for job title
        for pattern in title_patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                if current_entry:
                    experience.append(current_entry)
                
                # Extract dates if present
                date_match = re.search(r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})\s*[-–]\s*(Present|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4})', line, re.IGNORECASE)
                
                current_entry = {
                    'title': match.group(0).strip(),
                    'dates': date_match.group(0) if date_match else None,
                    'raw': line[:100]
                }
                break
    
    if current_entry:
        experience.append(current_entry)
    
    return experience


def _extract_projects(text: str) -> List[Dict[str, str]]:
    """Extract project information"""
    projects = []
    
    # Look for GitHub links associated with projects
    github_pattern = r'(?:https?://)?github\.com/[A-Za-z0-9_-]+/([A-Za-z0-9_-]+)'
    github_matches = re.findall(github_pattern, text)
    
    for repo in github_matches[:5]:  # Limit to 5 projects
        projects.append({
            'name': repo.replace('-', ' ').replace('_', ' ').title(),
            'url': f'https://github.com/{repo}'
        })
    
    return projects


def _validate_extraction(result: Dict[str, Any]) -> Dict[str, Any]:
    """Validate the extraction results"""
    validation = {
        'valid': True,
        'fields_found': [],
        'fields_missing': [],
        'confidence': 0,
        'warnings': [],
        'summary': {}
    }
    
    # Required fields
    required_fields = ['email', 'name']
    optional_fields = ['phone', 'github', 'linkedin']
    skill_fields = ['programming_languages', 'frameworks', 'skills']
    
    # Check required fields
    for field in required_fields:
        if result.get(field):
            validation['fields_found'].append(field)
        else:
            validation['fields_missing'].append(field)
            validation['warnings'].append(f'Could not extract {field}')
    
    # Check optional fields
    for field in optional_fields:
        if result.get(field):
            validation['fields_found'].append(field)
    
    # Check skill fields
    total_skills = 0
    for field in skill_fields:
        skills = result.get(field, [])
        if skills:
            validation['fields_found'].append(field)
            total_skills += len(skills)
    
    if total_skills == 0:
        validation['warnings'].append('No technical skills detected')
    
    # Calculate confidence score
    confidence = 0
    weights = {
        'email': 20,
        'name': 15,
        'phone': 10,
        'github': 10,
        'linkedin': 5,
        'programming_languages': 15,
        'frameworks': 10,
        'skills': 10,
        'experience': 5,
    }
    
    for field, weight in weights.items():
        value = result.get(field)
        if value and (not isinstance(value, list) or len(value) > 0):
            confidence += weight
    
    validation['confidence'] = min(confidence, 100)
    validation['valid'] = validation['confidence'] >= 30  # At least 30% confidence
    
    # Create summary for display
    validation['summary'] = {
        'contact_info': bool(result.get('email') or result.get('phone')),
        'has_name': bool(result.get('name')),
        'tech_skills_count': total_skills,
        'has_education': len(result.get('education', [])) > 0,
        'has_experience': len(result.get('experience', [])) > 0,
    }
    
    return validation


def extract_skills_from_text(text: str) -> List[str]:
    """Extract just skills from text"""
    result = _extract_info_enhanced(text)
    return result.get('all_technical_skills', [])


def validate_parsed_data(parsed_data: Dict[str, Any]) -> Dict[str, Any]:
    """Validate already parsed data"""
    return _validate_extraction(parsed_data)
